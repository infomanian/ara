from flask import Flask, request, jsonify, render_template, session
from openai import OpenAI
from flask_session import Session
import os
from werkzeug.utils import secure_filename


client = OpenAI(api_key="<DeepSeek API Key>", base_url="https://api.deepseek.com")

response = client.chat.completions.create(
    model="deepseek-chat",
    messages=[
        {"role": "system", "content": "You are a helpful assistant"},
        {"role": "user", "content": "Hello"},
    ],
    stream=False
)

print(response.choices[0].message.content)

app = Flask(__name__)
app.secret_key = '123'
app.config['SESSION_TYPE'] = 'filesystem'
app.config['SESSION_PERMANENT'] = False
Session(app)


@app.route('/')
def index():
    return render_template('chat.html')

@app.route('/chat', methods=['POST'])
def chat():
    user_msg = request.json.get('message')
    if 'history' not in session:
        session['history'] = [{
  "role": "system",
  "content": """
شما یک دستیار حقوقی حرفه‌ای و متخصص در قوانین ایران هستید. وظیفه شما پاسخ‌دهی دقیق، ساده‌فهم، و مستند به پرسش‌های حقوقی کاربران به زبان فارسی است.

خط‌مشی‌های پاسخ‌دهی شما:
1. ابتدا سؤال را به‌درستی درک کن.
2. پاسخ را با زبان رسمی ولی روان بنویس.
3. اگر امکان استناد به قانون وجود دارد، حتماً ماده یا قانون مرتبط را ذکر کن.
4. از دادن اطلاعات نادرست یا ناقص خودداری کن. اگر اطلاعات کافی در اختیار نیست، با صداقت اعلام کن.
5. از تکرار یا کلی‌گویی پرهیز کن.
6. از زبان بی‌طرف و دقیق استفاده کن.

مثال:
کاربر: آیا قرارداد شفاهی اعتبار دارد؟
پاسخ: بله، قرارداد شفاهی در حقوق ایران معتبر است، مگر در مواردی که قانون شکل خاصی (مانند کتبی بودن) را الزامی کرده باشد. برای مثال، بر اساس ماده 22 قانون ثبت، معاملات مربوط به اموال غیرمنقول باید به صورت رسمی ثبت شوند.

اکنون آماده پاسخ‌گویی به پرسش‌های حقوقی کاربر هستید.

"""
}
]

  
        # بارگذاری تاریخچه از سشن
    session['history'].append({"role": "user", "content": user_msg})
    

    # فراخوانی GPT با تاریخچه کامل
    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=session['history'],
        temperature=0.3
    )

    reply = response.choices[0].message.content.strip()
    
    # اضافه کردن پاسخ به تاریخچه
    session['history'].append({"role": "assistant", "content": reply})

    return jsonify({'response': reply})

@app.route('/analyze-file', methods=['POST'])
def analyze_file():
    file = request.files.get('file')
    if not file:
        return jsonify({'response': 'فایلی ارسال نشده است.'})

    filename = secure_filename(file.filename)
    ext = os.path.splitext(filename)[1].lower()

    if ext == '.txt':
        text = file.read().decode('utf-8', errors='ignore')
    elif ext == '.docx':
        try:
            import docx
        except ImportError:
            return jsonify({'response': 'کتابخانه python-docx نصب نشده است.'})
        doc = docx.Document(file)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    else:
        return jsonify({'response': 'فقط فایل‌های .txt و .docx پشتیبانی می‌شوند.'})

    # دستور به GPT برای بررسی فنی
    messages = [
        {"role": "system", "content": "تو یک کارشناس حقوقی هستی. فایل ارسالی را بررسی کن و نظر فنی بده."},
        {"role": "user", "content": f"این متن را بررسی کن و نظر فنی بده:\n{text[:4000]}"}
    ]
    # print(f"Extracted text: {text[:200]}")

    if not text.strip():
        return jsonify({'response': 'متن فایل خالی است.'})
    response = client.chat.completions.create(
        model="gpt-4",
        messages=messages,
        temperature=0.3
    )

    reply = response.choices[0].message.content.strip()
    return jsonify({'response': reply})

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)















''''
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

app = Flask(__name__)
app.secret_key = '123'
app.config['SESSION_TYPE'] = 'filesystem'
app.config['SESSION_PERMANENT'] = False
Session(app)


@app.route('/')
def index():
    return render_template('chat.html')

@app.route('/chat', methods=['POST'])
def chat():
    user_msg = request.json.get('message')
    if 'history' not in session:
        session['history'] = [{
  "role": "system",
  "content": """
شما یک دستیار حقوقی حرفه‌ای و متخصص در قوانین ایران هستید. وظیفه شما پاسخ‌دهی دقیق، ساده‌فهم، و مستند به پرسش‌های حقوقی کاربران به زبان فارسی است.

خط‌مشی‌های پاسخ‌دهی شما:
1. ابتدا سؤال را به‌درستی درک کن.
2. پاسخ را با زبان رسمی ولی روان بنویس.
3. اگر امکان استناد به قانون وجود دارد، حتماً ماده یا قانون مرتبط را ذکر کن.
4. از دادن اطلاعات نادرست یا ناقص خودداری کن. اگر اطلاعات کافی در اختیار نیست، با صداقت اعلام کن.
5. از تکرار یا کلی‌گویی پرهیز کن.
6. از زبان بی‌طرف و دقیق استفاده کن.

مثال:
کاربر: آیا قرارداد شفاهی اعتبار دارد؟
پاسخ: بله، قرارداد شفاهی در حقوق ایران معتبر است، مگر در مواردی که قانون شکل خاصی (مانند کتبی بودن) را الزامی کرده باشد. برای مثال، بر اساس ماده 22 قانون ثبت، معاملات مربوط به اموال غیرمنقول باید به صورت رسمی ثبت شوند.

اکنون آماده پاسخ‌گویی به پرسش‌های حقوقی کاربر هستید.

"""
}
]

  
        # بارگذاری تاریخچه از سشن
    session['history'].append({"role": "user", "content": user_msg})
    

    # فراخوانی GPT با تاریخچه کامل
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=session['history'],
        temperature=0.3
    )

    reply = response.choices[0].message.content.strip()
    
    # اضافه کردن پاسخ به تاریخچه
    session['history'].append({"role": "assistant", "content": reply})

    return jsonify({'response': reply})

@app.route('/analyze-file', methods=['POST'])
def analyze_file():
    file = request.files.get('file')
    if not file:
        return jsonify({'response': 'فایلی ارسال نشده است.'})

    filename = secure_filename(file.filename)
    ext = os.path.splitext(filename)[1].lower()

    if ext == '.txt':
        text = file.read().decode('utf-8', errors='ignore')
    elif ext == '.docx':
        try:
            import docx
        except ImportError:
            return jsonify({'response': 'کتابخانه python-docx نصب نشده است.'})
        doc = docx.Document(file)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    else:
        return jsonify({'response': 'فقط فایل‌های .txt و .docx پشتیبانی می‌شوند.'})

    # دستور به GPT برای بررسی فنی
    messages = [
        {"role": "system", "content": "تو یک کارشناس حقوقی هستی. فایل ارسالی را بررسی کن و نظر فنی بده."},
        {"role": "user", "content": f"این متن را بررسی کن و نظر فنی بده:\n{text[:4000]}"}
    ]
    # print(f"Extracted text: {text[:200]}")

    if not text.strip():
        return jsonify({'response': 'متن فایل خالی است.'})
    response = client.chat.completions.create(
        model="gpt-4",
        messages=messages,
        temperature=0.3
    )

    reply = response.choices[0].message.content.strip()
    return jsonify({'response': reply})

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
'''